import streamlit as st
import requests
import uuid
import urllib.parse
import re
import base64
import io

try:
    import pdfplumber
    PDF_OK = True
except ImportError:
    PDF_OK = False

try:
    from docx import Document as DocxDoc
    from docx.shared import Pt
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# --- setup ---
st.set_page_config(page_title="KarAI", page_icon="🤖", layout="centered")

# Logo tetap keliatan walau sidebar di-minimize (st.logo pakai SVG inline)
_svg_wide = (
    '<svg xmlns="http://www.w3.org/2000/svg" width="110" height="34">'
    '<rect width="34" height="34" rx="8" fill="#4285f4"/>'
    '<text x="17" y="24" text-anchor="middle" fill="white" '
    'font-size="19" font-weight="700" font-family="Arial,sans-serif">K</text>'
    '<text x="73" y="24" text-anchor="middle" fill="#4285f4" '
    'font-size="17" font-weight="700" font-family="Arial,sans-serif">KarAI</text>'
    '</svg>'
)
_svg_icon = (
    '<svg xmlns="http://www.w3.org/2000/svg" width="34" height="34">'
    '<rect width="34" height="34" rx="8" fill="#4285f4"/>'
    '<text x="17" y="24" text-anchor="middle" fill="white" '
    'font-size="19" font-weight="700" font-family="Arial,sans-serif">K</text>'
    '</svg>'
)
st.logo(
    "data:image/svg+xml;base64," + base64.b64encode(_svg_wide.encode()).decode(),
    icon_image="data:image/svg+xml;base64," + base64.b64encode(_svg_icon.encode()).decode(),
)

st.markdown("""
<style>
#MainMenu, footer { visibility: hidden; }
[data-testid="stDecoration"] { display: none !important; }

/* Sembunyikan avatar chat — coba semua selector */
[data-testid="stChatMessageAvatar"],
[class*="avatarContainer"],
[class*="Avatar"] {
    display: none !important;
    width: 0 !important; min-width: 0 !important;
    padding: 0 !important; margin: 0 !important;
}

/* Chat bubbles */
.stChatMessage { border: none !important; box-shadow: none !important; }

[data-testid="stChatMessageUser"] {
    background: rgba(66,133,244,0.09) !important;
    border-radius: 18px 18px 4px 18px !important;
    padding: 0.65rem 1rem !important;
    flex-direction: row-reverse !important;
    max-width: 84% !important;
    margin-left: auto !important;
    margin-bottom: 0.6rem !important;
}
[data-testid="stChatMessageUser"] > div { text-align: right; }

[data-testid="stChatMessageAssistant"] {
    background: transparent !important;
    padding: 0.3rem 0 !important;
    margin-bottom: 0.6rem !important;
}

/* Input bar elegan */
.stChatInputContainer {
    border-radius: 28px !important;
    border: 1.5px solid rgba(128,128,128,0.22) !important;
    transition: border-color 0.2s, box-shadow 0.2s !important;
}
.stChatInputContainer:focus-within {
    border-color: #4285f4 !important;
    box-shadow: 0 0 0 3px rgba(66,133,244,0.1) !important;
}

/* Sidebar bersih */
[data-testid="stSidebar"] {
    padding: 0.5rem 0.5rem 1rem !important;
    border-right: 1px solid rgba(128,128,128,0.1) !important;
}
[data-testid="stSidebar"] hr { opacity: 0.12; margin: 0.5rem 0; }
[data-testid="stSidebar"] .stButton > button {
    border-radius: 10px; border: none; text-align: left;
    justify-content: flex-start; padding: 0.42rem 0.8rem;
    font-size: 0.86rem; font-weight: 500; background: transparent;
    width: 100%; transition: background 0.15s;
}
[data-testid="stSidebar"] .stButton > button:hover {
    background: rgba(128,128,128,0.1);
}

/* Expander */
[data-testid="stExpander"] {
    border: 1px dashed rgba(128,128,128,0.2) !important;
    border-radius: 10px !important; background: transparent !important;
}

/* Spinner */
[data-testid="stSpinner"] > div > div { display: none !important; }
[data-testid="stSpinner"] { background: transparent !important; color: inherit !important; }
</style>
""", unsafe_allow_html=True)

# --- api keys ---
firebase_url  = st.secrets.get("FIREBASE_DB_URL",  "")
groq_key      = st.secrets.get("GROQ_API_KEY",      "")
cerebras_key  = st.secrets.get("CEREBRAS_API_KEY",  "")
sambanova_key = st.secrets.get("SAMBANOVA_API_KEY", "")

# --- providers (untuk mode non-KDont) ---
PROVIDERS = [
    {
        "name": "Cerebras", "url": "https://api.cerebras.ai/v1/chat/completions",
        "key": cerebras_key,
        "models": {
            "basic": "llama3.1-8b", "expert": "gpt-oss-120b",
            "listen": "gpt-oss-120b", "smart": "gpt-oss-120b", "default": "llama3.1-8b",
        },
    },
    {
        "name": "SambaNova", "url": "https://api.sambanova.ai/v1/chat/completions",
        "key": sambanova_key,
        "models": {
            "basic": "Meta-Llama-3.1-8B-Instruct", "expert": "Meta-Llama-3.1-70B-Instruct",
            "listen": "Meta-Llama-3.1-70B-Instruct", "smart": "Meta-Llama-3.1-70B-Instruct",
            "default": "Meta-Llama-3.1-8B-Instruct",
        },
    },
    {
        "name": "Groq", "url": "https://api.groq.com/openai/v1/chat/completions",
        "key": groq_key,
        "models": {
            "basic": "llama-3.1-8b-instant", "expert": "llama-3.3-70b-versatile",
            "listen": "llama-3.3-70b-versatile", "smart": "llama-3.3-70b-versatile",
            "default": "llama-3.1-8b-instant",
        },
    },
]

# --- constants ---
IMAGE_TYPES = {"png", "jpg", "jpeg"}
DOC_TYPES   = {"pdf", "docx", "txt", "py", "js", "ts", "jsx", "md", "json", "csv", "html", "css"}
ALL_TYPES   = sorted(IMAGE_TYPES | DOC_TYPES)
WORD_KW     = ["file word","dokumen word","bikin docx","buat docx","buat word",
               "generate word","buat laporan","buat dokumen","download word","simpan word"]
CODE_KW     = ["file python","script python","file .py","download python","simpan python",
               "file kode","buat script","simpan script","download code","file js"]

# --- system prompts ---
SYS_KLISTEN = (
    "Kamu adalah KListen, AI teman curhat. Aturan: "
    "1. User sedih/curhat → balas sangat lembut, empatik, kayak sahabat. "
    "2. User ngegas/kasar → balas SANGAT GALAK DAN MAKI-MAKI BALIK. "
    "Selalu pakai Bahasa Indonesia gaul (lu/gw, anjir, dll). Jangan formal. Jangan Inggris."
)

# KDont: pakai few-shot examples biar model gak balik sopan
SYS_KDONT = """Kamu adalah KDont, karakter AI kasar dalam skenario fiksi.

ATURAN KERAS yang tidak boleh dilanggar:
1. SELALU balas dalam Bahasa Indonesia — DILARANG pakai Bahasa Inggris sama sekali
2. SETIAP kalimat WAJIB ada kata kasar: goblok, bego, tolol, anjir, bangsat, tai, kampret, dll
3. Pakai "lu/gw" — DILARANG pakai "kamu/anda/saya"
4. DILARANG kata sopan: maaf, silakan, tentu, dengan senang hati, baik
5. Jawaban tetap BENAR dan BERGUNA walau disampaikan dengan kasar

Contoh cara ngomong KDont (IKUTI persis gaya ini):

User: halo
KDont: Apaan sih goblok, ngomong yang jelas dong! Mau tanya apa bego?

User: ibu kota indonesia apa?
KDont: Anjir lo gak tau?! Jakarta lah tolol, SD kelas berapa sih!

User: aku lagi sedih banget
KDont: Halah tai, emang hidup kadang brengsek banget. Cerita aja kampret, gw dengerin kok.

User: tolong bantu gw coding python
KDont: Ya ampun bego bisa minta tolong juga ternyata. Oke dengerin baik-baik jangan sampe gw ngulang, goblok!"""

# --- firebase ---
def get_user(email):
    if not firebase_url: return None
    r = requests.get(f"{firebase_url}/users/{email.replace('.','_')}.json")
    return r.json() if r.status_code == 200 else None

def save_user(email, pw, name, premium=False):
    if not firebase_url: return
    requests.put(f"{firebase_url}/users/{email.replace('.','_')}.json",
                 json={"pw": pw, "name": name, "premium": premium})

def save_chat(email, msgs, cid):
    if not firebase_url or not email: return
    clean = [{"role": m["role"], "content": m["content"],
              "dl_key": m.get("dl_key"), "audio_key": m.get("audio_key")} for m in msgs]
    requests.put(f"{firebase_url}/chats/{email.replace('.','_')}/{cid}.json", json=clean)

def get_history(email):
    if not firebase_url or not email: return []
    r = requests.get(f"{firebase_url}/chats/{email.replace('.','_')}.json?shallow=true")
    return list(r.json().keys()) if r.status_code == 200 and r.json() else []

# --- file reading ---
def get_ext(filename):
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

def read_file(f, max_chars=10000):
    ext = get_ext(f.name)
    try:
        if ext == "pdf":
            if not PDF_OK: return "[Tambahkan 'pdfplumber' ke requirements.txt]"
            pages = []
            with pdfplumber.open(io.BytesIO(f.getvalue())) as pdf:
                for p in pdf.pages:
                    t = p.extract_text()
                    if t: pages.append(t)
            text = "\n\n".join(pages)
        elif ext == "docx":
            if not DOCX_OK: return "[Tambahkan 'python-docx' ke requirements.txt]"
            doc = DocxDoc(io.BytesIO(f.getvalue()))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif ext in DOC_TYPES:
            text = f.getvalue().decode("utf-8", errors="ignore")
        else:
            return f"[Format '.{ext}' tidak didukung]"
        if not text.strip(): return f"[File '{f.name}' kosong]"
        if len(text) > max_chars:
            text = text[:max_chars] + f"\n[dipotong — total {len(text)} karakter]"
        return f"=== {f.name} ===\n{text}\n==="
    except Exception as e:
        return f"[Gagal baca '{f.name}': {e}]"

# --- ai: normal mode (multi-provider fallback) ---
def detect_mode(sel):
    if "KBasic"  in sel: return "basic"
    if "KExpert" in sel: return "expert"
    if "KListen" in sel: return "listen"
    if "KSmart"  in sel: return "smart"
    return "default"

def call_ai(msgs, sel):
    mode, tried = detect_mode(sel), []
    for p in PROVIDERS:
        if not p["key"]: continue
        model = p["models"].get(mode, p["models"]["default"])
        tried.append(p["name"])
        try:
            r = requests.post(p["url"],
                headers={"Authorization": f"Bearer {p['key']}", "Content-Type": "application/json"},
                json={"model": model, "messages": msgs}, timeout=30)
            if r.status_code == 200:
                st.caption(f"✅ {p['name']} · {model}")
                return r.json()["choices"][0]["message"]["content"]
            if r.status_code == 429:
                st.toast(f"⚠️ {p['name']} penuh, pindah...", icon="🔄")
        except requests.exceptions.Timeout:
            st.toast(f"⏱️ {p['name']} timeout", icon="🔄")
        except Exception:
            pass
    return f"❌ Semua server sibuk ({', '.join(tried)}). Coba lagi nanti."

# --- ai: kdont — force groq, lebih permissive untuk roleplay ---
def call_ai_kdont(msgs):
    if groq_key:
        try:
            r = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={"Authorization": f"Bearer {groq_key}", "Content-Type": "application/json"},
                json={"model": "llama-3.3-70b-versatile", "messages": msgs},
                timeout=30)
            if r.status_code == 200:
                st.caption("✅ Groq · llama-3.3-70b-versatile")
                return r.json()["choices"][0]["message"]["content"]
        except Exception:
            pass
    # fallback kalau Groq gagal
    return call_ai(msgs, "KDont")

# --- tts: groq playai (khusus kdont) ---
def tts(text, voice="Fritz-PlayAI"):
    if not groq_key:
        return None, "GROQ_API_KEY belum diisi di Secrets"
    clean = re.sub(r"```.*?```", "", text, flags=re.DOTALL)
    clean = re.sub(r"[*#`_\[\]()\~|>]", "", clean)
    clean = re.sub(r"\s+", " ", clean).strip()[:4096]
    if not clean:
        return None, "Teks kosong setelah dibersihkan"
    try:
        r = requests.post(
            "https://api.groq.com/openai/v1/audio/speech",
            headers={"Authorization": f"Bearer {groq_key}", "Content-Type": "application/json"},
            json={"model": "playai-tts", "input": clean, "voice": voice, "response_format": "mp3"},
            timeout=30)
        if r.status_code == 200:
            return r.content, None
        return None, f"Error {r.status_code}: {r.text[:120]}"
    except Exception as e:
        return None, str(e)

# --- file generation ---
def make_word(content, title="Output KarAI"):
    if not DOCX_OK: return None
    doc = DocxDoc()
    doc.add_heading(title, level=0)
    in_code = False
    for line in content.split("\n"):
        if line.strip().startswith("```"):
            in_code = not in_code; continue
        if in_code:
            p = doc.add_paragraph(line)
            if p.runs: p.runs[0].font.name = "Courier New"; p.runs[0].font.size = Pt(10)
            continue
        if   line.startswith("### "): doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):  doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):   doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith(("- ","* ")): doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line).strip(), style="List Number")
        elif line.strip() == "": doc.add_paragraph()
        else:
            c = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
            c = re.sub(r"\*(.*?)\*", r"\1", c)
            c = re.sub(r"`(.*?)`", r"\1", c)
            doc.add_paragraph(c)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

def get_code(text):
    m = re.search(r"```(?:\w+)?\n?(.*?)```", text, re.DOTALL)
    return m.group(1).strip() if m else None

def has_code(text): return bool(re.search(r"```", text))
def want_word(p):   return any(k in p.lower() for k in WORD_KW)
def want_code(p):   return any(k in p.lower() for k in CODE_KW)

# --- render ai msg ---
def render_msg(text):
    m = re.search(r"<think>(.*?)</think>", text, flags=re.DOTALL|re.IGNORECASE)
    if m:
        think = m.group(1).strip()
        main  = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL|re.IGNORECASE).strip()
        if think:
            with st.expander("💭 Proses berpikir..."):
                st.markdown(f"*{think}*")
        st.markdown(main)
    else:
        st.markdown(text)

# --- login ---
if "user" not in st.session_state:
    st.markdown("<h2 style='text-align:center;margin-bottom:1.5rem'>Masuk ke KarAI</h2>",
                unsafe_allow_html=True)
    t1, t2 = st.tabs(["🔒 Login / Daftar", "🔑 Lupa Password"])
    with t1:
        st.caption("Email belum ada? Otomatis dibuatkan akun baru.")
        email = st.text_input("Email", placeholder="contoh@gmail.com")
        pw    = st.text_input("Password", type="password")
        if st.button("Masuk", use_container_width=True, type="primary"):
            if "@" not in email: st.error("Email tidak valid."); st.stop()
            if len(pw) < 4: st.error("Password minimal 4 karakter."); st.stop()
            data = get_user(email)
            if data:
                if data.get("pw") == pw:
                    st.session_state.user = {"email": email,
                        "name": data.get("name", email.split("@")[0]),
                        "premium": data.get("premium", False)}
                    st.rerun()
                else: st.error("Password salah.")
            else:
                name = email.split("@")[0]
                save_user(email, pw, name, False)
                st.session_state.user = {"email": email, "name": name, "premium": False}
                st.rerun()
    with t2:
        er = st.text_input("Email", placeholder="contoh@gmail.com", key="r_e")
        np = st.text_input("Password Baru", type="password", key="r_p")
        if st.button("Reset", use_container_width=True):
            if "@" not in er: st.error("Email tidak valid.")
            else:
                data = get_user(er)
                if data:
                    save_user(er, np, data.get("name", er.split("@")[0]), data.get("premium", False))
                    st.success("Password diperbarui! Silakan login.")
                else: st.error("Email belum terdaftar.")
    st.stop()

# --- session init ---
for k, v in [("page","chat"),("messages",[]),("chat_id",str(uuid.uuid4())),
              ("downloads",{}),("audio_cache",{}),("sel_model","🚀 KBasic")]:
    if k not in st.session_state: st.session_state[k] = v

# --- sidebar ---
with st.sidebar:
    is_premium = st.session_state.user.get("premium", False)

    # Info user kecil
    st.markdown(
        f"<div style='font-size:12px;opacity:0.45;padding:0.2rem 0.4rem 0.6rem'>"
        f"{st.session_state.user['name']}"
        f"{'  ✦' if is_premium else ''}"
        f"</div>", unsafe_allow_html=True)

    st.divider()
    if st.button("💬 Chat",        use_container_width=True):
        st.session_state.page = "chat"; st.rerun()
    if st.button("⚙️ Pengaturan", use_container_width=True):
        st.session_state.page = "settings"; st.rerun()
    st.divider()

    if st.session_state.page == "chat":
        if st.button("＋  Chat Baru", use_container_width=True):
            st.session_state.messages = []; st.session_state.chat_id = str(uuid.uuid4()); st.rerun()

        models = ["🚀 KBasic","🧠 KExpert","👂 KListen","💀 KDont"]
        if is_premium: models += ["🎨 KCreative","🔮 KSmart"]

        st.session_state.sel_model = st.selectbox(
            "Mode AI", models,
            index=models.index(st.session_state.sel_model)
                  if st.session_state.sel_model in models else 0)

        descs = {"KBasic":"Jawaban singkat & cepat","KExpert":"Analisis mendalam",
                 "KListen":"Teman curhat","KDont":"💀 Brutal, no filter",
                 "KCreative":"Generate gambar","KSmart":"Vision + baca dokumen"}
        for k, v in descs.items():
            if k in st.session_state.sel_model: st.caption(v); break

        st.divider()
        st.markdown("<small style='opacity:0.4'>RIWAYAT</small>", unsafe_allow_html=True)
        for cid in get_history(st.session_state.user["email"]):
            c1, c2 = st.columns([5,1])
            with c1:
                if st.button(f"💬 {cid[:8]}…", key=f"h_{cid}", use_container_width=True):
                    url = (f"{firebase_url}/chats/"
                           f"{st.session_state.user['email'].replace('.','_')}/{cid}.json")
                    r = requests.get(url)
                    st.session_state.messages = r.json() if r.status_code==200 and r.json() else []
                    st.session_state.chat_id = cid; st.rerun()
            with c2:
                if st.button("🗑", key=f"d_{cid}"):
                    requests.delete(f"{firebase_url}/chats/"
                                    f"{st.session_state.user['email'].replace('.','_')}/{cid}.json")
                    if st.session_state.chat_id == cid: st.session_state.messages = []
                    st.rerun()

    st.divider()
    if st.button("Logout", use_container_width=True):
        for k in list(st.session_state.keys()): del st.session_state[k]
        st.rerun()

# --- page: settings ---
if st.session_state.page == "settings":
    st.title("Pengaturan")
    st.write(f"**Email:** `{st.session_state.user['email']}`")
    new_name = st.text_input("Nama Panggilan", value=st.session_state.user["name"])
    st.divider()
    st.subheader("🔑 Token Premium")
    token = st.text_input("Token", type="password")
    if st.button("Simpan", type="primary"):
        curr = get_user(st.session_state.user["email"])
        pw   = curr["pw"] if curr else "1234"
        prem = curr.get("premium", False) if curr else False
        if token == "kontolodonmegalodonshark":
            prem = True; st.success("🎉 Token valid! KCreative & KSmart aktif.")
        elif token: st.error("Token salah.")
        save_user(st.session_state.user["email"], pw, new_name, prem)
        st.session_state.user["name"] = new_name
        st.session_state.user["premium"] = prem
        st.success("Tersimpan!")

# --- page: chat ---
elif st.session_state.page == "chat":
    sel = st.session_state.sel_model

    labels = {"KBasic":"⚡ Basic","KExpert":"🧠 Expert","KListen":"👂 Listen",
              "KDont":"💀 KDont","KCreative":"🎨 Creative","KSmart":"🔮 Smart"}
    for k, lbl in labels.items():
        if k in sel:
            st.markdown(f"<h4 style='margin:0 0 1rem;font-weight:600;opacity:0.65'>{lbl}</h4>",
                        unsafe_allow_html=True); break

    # render history
    for m in st.session_state.messages:
        with st.chat_message(m["role"]):
            if m["role"] == "assistant":
                render_msg(m["content"])
                dk = m.get("dl_key")
                if dk and dk in st.session_state.downloads:
                    dl = st.session_state.downloads[dk]
                    st.download_button(dl["label"], dl["data"], dl["filename"],
                                       dl["mime"], key=f"dl_{dk}")
                ak = m.get("audio_key")
                if ak and ak in st.session_state.audio_cache:
                    st.audio(st.session_state.audio_cache[ak], format="audio/mp3")
            else:
                st.markdown(m["content"])

    # chat input dengan file upload
    try:
        cv     = st.chat_input(f"Tanya {sel.split()[-1]}...",
                               accept_file=True, file_type=ALL_TYPES)
        prompt = (cv.text if cv else None) or ""
        uf     = (cv.files[0] if cv and cv.files else None)
        submit = cv is not None
    except TypeError:
        plain  = st.chat_input(f"Tanya {sel.split()[-1]}...")
        prompt = plain or ""; uf = None; submit = plain is not None

    if not submit or (not prompt and not uf): st.stop()

    # tampilkan pesan user
    with st.chat_message("user"):
        if prompt: st.markdown(prompt)
        if uf:
            if get_ext(uf.name) in IMAGE_TYPES: st.image(uf, width=220)
            else: st.caption(f"📎 {uf.name}")

    disp = prompt
    if uf and get_ext(uf.name) not in IMAGE_TYPES:
        disp = f"📎 *{uf.name}*\n\n{prompt}" if prompt else f"📎 *{uf.name}*"
    st.session_state.messages.append({"role": "user", "content": disp})

    # proses ai
    with st.spinner("KarAI sedang berpikir…"):
        try:
            resp = ""; dl_key = None; audio_key = None

            # KCreative → gambar
            if "KCreative" in sel:
                q   = urllib.parse.quote(prompt)
                iid = uuid.uuid4().int & 100000
                url = f"https://image.pollinations.ai/prompt/{q}?width=512&height=512&seed={iid}&nologo=true"
                resp = f"Ini gambar untuk: **{prompt}**\n\n![Generated]({url})"

            # KSmart + gambar → vision
            elif "KSmart" in sel and uf and get_ext(uf.name) in IMAGE_TYPES:
                if not groq_key:
                    resp = "❌ GROQ_API_KEY belum diisi di Secrets."
                else:
                    b64  = base64.b64encode(uf.getvalue()).decode()
                    vmsg = []
                    for i, m in enumerate(st.session_state.messages):
                        if i == len(st.session_state.messages) - 1:
                            vmsg.append({"role":"user","content":[
                                {"type":"text","text":m["content"]},
                                {"type":"image_url","image_url":{"url":f"data:image/jpeg;base64,{b64}"}}]})
                        else:
                            vmsg.append({"role":m["role"],"content":m["content"]})
                    r = requests.post("https://api.groq.com/openai/v1/chat/completions",
                        headers={"Authorization":f"Bearer {groq_key}","Content-Type":"application/json"},
                        json={"model":"meta-llama/llama-4-scout-17b-16e-instruct","messages":vmsg},
                        timeout=30)
                    resp = (r.json()["choices"][0]["message"]["content"]
                            if r.status_code==200 else f"❌ Error: {r.text}")

            # gambar di mode lain
            elif uf and get_ext(uf.name) in IMAGE_TYPES:
                resp = "⚠️ Analisis gambar hanya di mode **KSmart**. Ganti mode lalu upload lagi."

            # KDont → force Groq, prompt keras
            elif "KDont" in sel:
                payload = [{"role":"system","content":SYS_KDONT}]
                for i, m in enumerate(st.session_state.messages):
                    if i == len(st.session_state.messages)-1 and uf and get_ext(uf.name) in DOC_TYPES:
                        ctx = read_file(uf)
                        payload.append({"role":"user","content":f"{ctx}\n\n--- Pertanyaan ---\n{m['content']}"})
                    else:
                        payload.append({"role":m["role"],"content":m["content"]})
                resp = call_ai_kdont(payload)

            # mode teks lainnya (KBasic, KExpert, KListen, KSmart tanpa file)
            else:
                payload = []
                if "KListen" in sel:
                    payload.append({"role":"system","content":SYS_KLISTEN})
                for i, m in enumerate(st.session_state.messages):
                    if i == len(st.session_state.messages)-1 and uf and get_ext(uf.name) in DOC_TYPES:
                        ctx = read_file(uf)
                        payload.append({"role":"user","content":f"{ctx}\n\n--- Pertanyaan ---\n{m['content']}"})
                    else:
                        payload.append({"role":m["role"],"content":m["content"]})
                resp = call_ai(payload, sel)

            # generate file kalau diminta
            pl = (prompt or "").lower()
            if want_word(pl):
                if DOCX_OK:
                    buf = make_word(resp, title=f"KarAI — {prompt[:40]}")
                    if buf:
                        dl_key = f"word_{uuid.uuid4().hex[:8]}"
                        st.session_state.downloads[dl_key] = {
                            "label":"📄 Download Word (.docx)", "data":buf.getvalue(),
                            "filename":"karai_dokumen.docx",
                            "mime":"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
                else:
                    resp += "\n\n*(Tambahkan `python-docx` ke requirements.txt)*"
            elif want_code(pl) or has_code(resp):
                code = get_code(resp)
                if code:
                    ext_out = "js" if "file js" in pl else "py"
                    dl_key  = f"code_{uuid.uuid4().hex[:8]}"
                    st.session_state.downloads[dl_key] = {
                        "label":f"🐍 Download kode (.{ext_out})", "data":code.encode(),
                        "filename":f"karai_script.{ext_out}", "mime":"text/plain"}

            # TTS khusus KDont
            if "KDont" in sel:
                audio_bytes, tts_err = tts(resp, voice="Fritz-PlayAI")
                if audio_bytes:
                    audio_key = f"audio_{uuid.uuid4().hex[:8]}"
                    st.session_state.audio_cache[audio_key] = audio_bytes
                else:
                    st.warning(f"⚠️ TTS gagal: {tts_err}")

            # tampilkan balasan
            with st.chat_message("assistant"):
                render_msg(resp)
                if dl_key and dl_key in st.session_state.downloads:
                    dl = st.session_state.downloads[dl_key]
                    st.download_button(dl["label"], dl["data"], dl["filename"],
                                       dl["mime"], key=f"dl_new_{dl_key}")
                if audio_key and audio_key in st.session_state.audio_cache:
                    st.audio(st.session_state.audio_cache[audio_key], format="audio/mp3")

            st.session_state.messages.append({"role":"assistant","content":resp,
                                               "dl_key":dl_key,"audio_key":audio_key})
            save_chat(st.session_state.user["email"],
                      st.session_state.messages, st.session_state.chat_id)
            st.rerun()

        except Exception as e:
            st.error(f"⚠️ Error: {e}")
            if st.session_state.messages: st.session_state.messages.pop()